from pptx import Presentation
from pptx.chart.data import CategoryChartData
from pptx.enum.chart import XL_CHART_TYPE
from pptx.util import Inches
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import ncap_standards as ncap_stds
import os

class Export2Document():

    def __init__(self, NCAP_data):
        if isinstance(NCAP_data, ncap_stds.NCAP_Standard):
            self.ncap_data = NCAP_data
        else:
            raise AttributeError("Export2PPTX object required 'NCAP_Standard' class object for creation")



    def brake_model_documentation(self):

        braketext = """
        Values used in simulation for Brake Model:
            1) t_dead = 100 ms
            2) t_ramp = 300 ms
            3) Decel_max = -10 m/s^2

        Note: This model is applied the same for all test cases.
              If the actual deceleration used in car is less than -10 m/s^2, then, the points for the test case might 
              change.
        """

        monte_carlo_assumptions = [
            "The impact speed calculation is done only for longitudinal motion. It has not yet been adopted for "
            "turning scenarios. Hence, the CCLTA scenario has not been included yet.",
            "The impact speed calculation does not take into account the motion of bicycle and pedestrian.",
            "If a testcase is not available in the test report of a scenario, then, the testcase gets awarded 0 points",
            "The entry/minimum test speed conditions are checked for all car-to-car cases, but not checked for "
            "Vru scenarios",
            "The initial distance between the ego and target is calculated as (measured TTC values X test speed) "
            "instead of the reading from RTE values.",
            "The impact speed calculation for CCRb scenario is not done, because, the measured TTC values already "
            "considers the deceleration of target. Hence, the initial distance estimate between ego and target will"
            " not be correct."
            ]

        # Write to HTML Report Brake model
        self.document.add_heading("Brake model", level=1)
        self.document.add_paragraph('Brake Model used for calculating impact speed from measured TTC values')
        self.document.add_picture(os.path.join("./images", "Brake_Model.png"), width=Inches(5))
        last_paragraph = self.document.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # report.AddImage(cHTMLReportImage(os.path.join(rp, "Brake_Model.png"), 'Scenario Median Scores', 1024, 800))
        # report.AddText('\n' + braketext)

        self.document.add_heading("Assumptions and Limitations in Monte-Carlo Simulations", level=2)
        # add a data row for each item
        for entry in monte_carlo_assumptions:
            self.document.add_paragraph(entry, style='List Bullet')

    def test_case_documentation(self):

        self.document.add_heading("Test Case Details", level=1)
        self.document.add_heading("Input Filename", level=2)
        for filename in self.ncap_data.input_files:
            self.document.add_paragraph(filename, style='List Bullet')

        # Add table for test case numbers
        self.document.add_heading("Test Case Details", level=2)
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run('Scenario').bold = True
        hdr_cells[1].paragraphs[0].add_run('No. of Test Cases').bold = True
        for sc, sc_obj in self.ncap_data.list_of_scenarios.items():
            row_cells = table.add_row().cells
            row_cells[0].text = sc
            row_cells[1].text = str(len(sc_obj.TEST_CASES))

        # Add missing scenarios
        self.document.add_heading("Missing Test cases", level=2)
        self.document.add_paragraph("Following scenarios were found in report but not included in "
                                    "scenario list for" + self.ncap_data.name + "standard")
        for scenario in self.ncap_data.missing_scenarios:
            self.document.add_paragraph(scenario, style='List Bullet')

        # Add warnings section
        self.document.add_heading("Warnings and Errors", level=2)
        for sc, sc_obj in self.ncap_data.list_of_scenarios.items():
            for warning in sc_obj.warning_list:
                self.document.add_paragraph(warning, style='List Bullet')

    def export_data(self):

        # Create document
        self.document = Document()
        self.document.add_heading(self.ncap_data.name + " Point Calculation Report", 0)

        # Add Brake model data
        self.brake_model_documentation()
        # Add Test case details
        self.test_case_documentation()

        self.document.save('demo.docx')

